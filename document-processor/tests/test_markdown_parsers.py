import sys
import tempfile
from pathlib import Path

import types

from docx import Document

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

import parser_any  # noqa: E402
from parser_docx import parse_docx  # noqa: E402


def test_docx_headings_are_rendered_as_markdown():
    with tempfile.TemporaryDirectory() as tmpdir:
        path = Path(tmpdir) / "sample.docx"
        doc = Document()
        doc.add_heading("Section title", level=2)
        doc.add_paragraph("Paragraph text")
        doc.save(path)

        blocks, _ = parse_docx(str(path), "doc-1")

        assert blocks[0]["text"].startswith("## Section title")
        assert blocks[0]["type"] == "heading"
        assert blocks[1]["text"] == "Paragraph text"


def test_pdf_pages_are_marked_as_markdown_headings(monkeypatch):
    class _StubPage:
        def __init__(self, text):
            self._text = text

        def extract_text(self, x_tolerance=2, y_tolerance=2):  # noqa: ARG002
            return self._text

    class _StubPDF:
        metadata = {}

        def __init__(self):
            self.pages = [_StubPage("Page body text")]

        def __enter__(self):
            return self

        def __exit__(self, exc_type, exc, tb):  # noqa: ANN001, ANN201
            return False

    def _stub_open(path):  # noqa: ARG001
        return _StubPDF()

    monkeypatch.setattr(parser_any, "pdfplumber", types.SimpleNamespace(open=_stub_open))

    blocks, _ = parser_any.parse_pdf("dummy.pdf", "doc-1")

    assert blocks[0]["text"].startswith("## Page 1")
    assert "Page body text" in blocks[0]["text"]
